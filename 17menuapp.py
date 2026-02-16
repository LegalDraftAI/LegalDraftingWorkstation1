import streamlit as st
import os
import pandas as pd
from datetime import datetime
from google import genai
from docx import Document
from fpdf import FPDF
import io
import urllib.parse

# 1. INITIALIZATION & DIRECTORY SETUP
VAULT_PATH = "private_vault"
HISTORY_FILE = "kerala_legal_history.csv"
KEY_FILE = "geminiapikey.py"
if not os.path.exists(VAULT_PATH): os.makedirs(VAULT_PATH)

st.set_page_config(page_title="Kerala Senior Advocate Workstation", layout="wide")

if "final_master" not in st.session_state: st.session_state.final_master = ""
if "search_results" not in st.session_state: st.session_state.search_results = []
if "widget_key" not in st.session_state: st.session_state.widget_key = 0

# 2. CORE UTILITY FUNCTIONS
def get_key():
    try:
        import geminiapikey
        return geminiapikey.GOOGLE_API_KEY
    except: return ""

def load_history():
    if os.path.exists(HISTORY_FILE):
        try: return pd.read_csv(HISTORY_FILE).tail(10).iloc[::-1]
        except: return pd.DataFrame()
    return pd.DataFrame()

def save_draft_to_history(dtype, content):
    new_data = pd.DataFrame([{"Date": datetime.now().strftime("%Y-%m-%d %H:%M"), "Type": dtype, "Draft": content}])
    new_data.to_csv(HISTORY_FILE, mode='a', header=not os.path.exists(HISTORY_FILE), index=False)

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    # Clean text to avoid latin-1 encoding crashes common in FPDF
    clean_text = text.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 8, clean_text)
    return pdf.output(dest='S').encode('latin-1')

def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. INTERFACE TABS
tab_work, tab_settings = st.tabs(["⚖️ Legal Workstation", "⚙️ Senior Admin Settings"])

# --- SETTINGS ---
with tab_settings:
    st.header("🔑 Senior Advocate API Configuration")
    stored_key = get_key()
    senior_key = st.text_input("Senior Advocate's Private Key:", value=stored_key, type="password")
    if st.button("💾 Save Settings"):
        with open(KEY_FILE, "w") as f: f.write(f'GOOGLE_API_KEY = "{senior_key}"')
        st.success("✅ Settings Saved!")
        st.rerun()

# --- WORKSTATION ---
with tab_work:
    USER_KEY = get_key()
    
    with st.sidebar:
        st.title("👨‍⚖️ Control Panel")
        if USER_KEY: st.success("🟢 Online")
        else: st.warning("🔴 Key Needed")

        # CASE SELECTION
        doc_type = st.selectbox("Petition Type", ["Bail Application", "NI Act (Sec 138)", "Writ Petition", "MC 125", "Domestic Violence", "MVOP", "Injunction", "Divorce"])
        jurisdiction = st.selectbox("Court Level", ["High Court", "District Court", "Family Court", "Munsiff Court"])
        is_hc = (jurisdiction == "High Court")
        selected_district = st.selectbox("District", ["Alappuzha", "Ernakulam", "Idukki", "Kannur", "Kasaragod", "Kollam", "Kottayam", "Kozhikode", "Malappuram", "Palakkad", "Pathanamthitta", "Thiruvananthapuram", "Thrissur", "Wayanad"], index=1 if is_hc else 0, disabled=is_hc)
        
        st.divider()
        uploaded = st.file_uploader("Style Vault Ref", type="docx")
        if uploaded:
            with open(os.path.join(VAULT_PATH, uploaded.name), "wb") as f: f.write(uploaded.getbuffer())
        vault_files = os.listdir(VAULT_PATH)
        selected_ref = st.selectbox("Mirror Reference:", ["None"] + vault_files)

        if st.button("🧹 Clear All", use_container_width=True):
            st.session_state.final_master = ""
            st.session_state.search_results = []
            st.session_state.widget_key += 1
            st.rerun()

        # HISTORY SECTION (STAYS AT BOTTOM)
        st.divider()
        st.subheader("📜 Recent History")
        hist_df = load_history()
        if not hist_df.empty:
            for i, row in hist_df.iterrows():
                if st.button(f"📄 {row['Type']} ({row['Date'][:10]})", key=f"h_{i}", use_container_width=True):
                    st.session_state.final_master = row['Draft']
                    st.rerun()

    st.title(f"Drafting: {doc_type}")
    facts = st.text_area("Case Facts:", height=150, key=f"f_{st.session_state.widget_key}")

    # COMMAND CENTER
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("🔍 Search API", use_container_width=True):
            st.session_state.search_results = [{"title": "Kerala HC Precedent 2025", "cite": "2025 KER 101", "sum": "Legal grounds confirmed."}]
    with c2:
        q = urllib.parse.quote_plus(facts if facts else "Kerala Law")
        st.link_button("🌐 Web Research", f"https://indiankanoon.org/search/?formInput={q}", use_container_width=True)
    with c3:
        if st.button("🚀 Standard Draft", use_container_width=True):
            if USER_KEY:
                with st.spinner("AI Drafting... (Live Timer)", show_time=True):
                    client = genai.Client(api_key=USER_KEY)
                    prompt = f"As a Kerala Lawyer, draft {doc_type} for {jurisdiction}. Facts: {facts}. STRICTLY USE 'PARTY A' for Petitioner and 'PARTY B' for Respondent. NO NAMES."
                    res = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    st.session_state.final_master = res.text
    with c4:
        if st.button("✨ Mirror Style", type="primary", use_container_width=True, disabled=(selected_ref=="None")):
            if USER_KEY:
                with st.spinner(f"Mirroring {selected_ref}...", show_time=True):
                    client = genai.Client(api_key=USER_KEY)
                    doc = Document(os.path.join(VAULT_PATH, selected_ref))
                    dna = "\n".join([p.text for p in doc.paragraphs[:15]])
                    prompt = f"MIMIC STYLE:\n{dna}\n\nTASK: Draft {doc_type} for {facts}. USE 'PARTY A' and 'PARTY B' for names."
                    res = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    st.session_state.final_master = res.text

    if st.session_state.search_results:
        with st.expander("📚 Related Precedents", expanded=True):
            for r in st.session_state.search_results: st.write(f"**{r['title']}** ({r['cite']})")

    # EDITOR & PARTY REPLACEMENT
    if st.session_state.final_master:
        st.divider()
        st.subheader("🛠️ Name Replacement Mapping")
        m1, m2 = st.columns(2)
        with m1:
            name_a = st.text_input("Name for PARTY A (Petitioner):")
            if st.button("Apply Petitioner Name"):
                st.session_state.final_master = st.session_state.final_master.replace("PARTY A", name_a)
                st.rerun()
        with m2:
            name_b = st.text_input("Name for PARTY B (Respondent):")
            if st.button("Apply Respondent Name"):
                st.session_state.final_master = st.session_state.final_master.replace("PARTY B", name_b)
                st.rerun()

        st.session_state.final_master = st.text_area("Live Editor:", value=st.session_state.final_master, height=500)
        
        # EXPORTS
        e1, e2, e3 = st.columns(3)
        with e1:
            if st.button("💾 Log to History", use_container_width=True):
                save_draft_to_history(doc_type, st.session_state.final_master)
                st.toast("Saved!")
        with e2:
            st.download_button("📥 MS Word", data=create_docx(st.session_state.final_master), file_name="draft.docx", use_container_width=True)
        with e3:
            st.download_button("📥 PDF", data=create_pdf(st.session_state.final_master), file_name="draft.pdf", use_container_width=True)