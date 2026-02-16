import streamlit as st
import os
import pandas as pd
from datetime import datetime
from google import genai
from docx import Document
from fpdf import FPDF
import io
import urllib.parse

# 1. INITIALIZATION
VAULT_PATH = "private_vault"
HISTORY_FILE = "kerala_legal_history.csv"
KEY_FILE = "geminiapikey.py"
if not os.path.exists(VAULT_PATH): os.makedirs(VAULT_PATH)

st.set_page_config(page_title="Kerala Senior Advocate Workstation", layout="wide")

if "final_master" not in st.session_state: st.session_state.final_master = ""
if "search_results" not in st.session_state: st.session_state.search_results = []
if "widget_key" not in st.session_state: st.session_state.widget_key = 0

# 2. CORE UTILITIES
def get_key():
    try:
        import geminiapikey
        return geminiapikey.GOOGLE_API_KEY
    except: return ""

def save_draft_to_history(dtype, content):
    new_data = pd.DataFrame([{"Date": datetime.now().strftime("%Y-%m-%d %H:%M"), "Type": dtype, "Draft": content}])
    new_data.to_csv(HISTORY_FILE, mode='a', header=not os.path.exists(HISTORY_FILE), index=False)

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Handling potential Unicode issues for PDF
    clean_text = text.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    return pdf.output(dest='S').encode('latin-1')

def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. TABS NAVIGATION
tab_work, tab_settings = st.tabs(["⚖️ Workstation", "⚙️ Admin Settings"])

with tab_settings:
    st.header("System Settings")
    input_key = st.text_input("Gemini API Key:", value=get_key(), type="password")
    if st.button("💾 Save & Initialize"):
        with open(KEY_FILE, "w") as f: f.write(f'GOOGLE_API_KEY = "{input_key}"')
        st.success("Key saved! App will use this on next run.")
        st.rerun()

with tab_work:
    USER_KEY = get_key()
    
    with st.sidebar:
        st.title("👨‍⚖️ Admin Panel")
        if USER_KEY: st.success("🟢 Connected")
        else: st.error("🔴 Key Missing")

        doc_type = st.selectbox("Petition Type", ["Bail Application", "NI Act", "Writ Petition", "Maintenance"])
        jurisdiction = st.selectbox("Court Level", ["High Court", "District Court"])
        is_hc = (jurisdiction == "High Court")
        districts = ["Alappuzha", "Ernakulam", "Idukki", "Kannur", "Kasaragod", "Kollam", "Kottayam", "Kozhikode", "Malappuram", "Palakkad", "Pathanamthitta", "Thiruvananthapuram", "Thrissur", "Wayanad"]
        selected_district = st.selectbox("District", districts, index=1 if is_hc else 0, disabled=is_hc)
        
        st.divider()
        uploaded = st.file_uploader("Vault: Add Ref Doc", type="docx")
        if uploaded:
            with open(os.path.join(VAULT_PATH, uploaded.name), "wb") as f: f.write(uploaded.getbuffer())
        
        vault_files = os.listdir(VAULT_PATH)
        selected_ref = st.selectbox("Style Reference:", ["None"] + vault_files)

        if st.button("🧹 Clear All", use_container_width=True):
            st.session_state.final_master = ""
            st.session_state.search_results = []
            st.session_state.widget_key += 1
            st.rerun()

    # MAIN WORK AREA
    st.title(f"Drafting: {doc_type}")
    facts = st.text_area("Case Facts:", height=150, key=f"f_{st.session_state.widget_key}")

    # COMMAND CENTER
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("🔍 Search API", use_container_width=True):
            st.session_state.search_results = [{"title": "State vs. Joy (2025)", "cite": "2025 KER 456", "sum": "Kerala High Court ruling on similar facts."}]
    with c2:
        q = urllib.parse.quote_plus(facts if facts else "Kerala Law")
        st.link_button("🌐 Web Research", f"https://indiankanoon.org/search/?formInput={q}", use_container_width=True)
    with c3:
        if st.button("🚀 Standard Draft", use_container_width=True):
            if USER_KEY:
                client = genai.Client(api_key=USER_KEY)
                res = client.models.generate_content(model='gemini-1.5-flash', contents=f"Draft {doc_type} for {selected_district}. Facts: {facts}")
                st.session_state.final_master = res.text
    with c4:
        if st.button("✨ Mirror Style", type="primary", use_container_width=True, disabled=(selected_ref=="None")):
            if USER_KEY:
                client = genai.Client(api_key=USER_KEY)
                doc = Document(os.path.join(VAULT_PATH, selected_ref))
                dna = "\n".join([p.text for p in doc.paragraphs[:15]])
                res = client.models.generate_content(model='gemini-1.5-flash', contents=f"STYLE: {dna}\n\nDraft {doc_type}: {facts}")
                st.session_state.final_master = res.text

    if st.session_state.search_results:
        with st.expander("📚 Related Precedents", expanded=True):
            for r in st.session_state.search_results:
                st.write(f"**{r['title']}** ({r['cite']}) - {r['sum']}")

    # EDITOR & EXPORTS
    if st.session_state.final_master:
        st.divider()
        st.subheader("📜 Live Editor")
        st.session_state.final_master = st.text_area("Edit Draft:", value=st.session_state.final_master, height=500)
        
        e1, e2, e3, e4 = st.columns(4)
        with e1:
            if st.button("💾 Save to History", use_container_width=True):
                save_draft_to_history(doc_type, st.session_state.final_master)
                st.toast("Saved!")
        with e2:
            st.download_button("📥 Download Word", data=create_docx(st.session_state.final_master), file_name="draft.docx", use_container_width=True)
        with e3:
            st.download_button("📥 Download PDF", data=create_pdf(st.session_state.final_master), file_name="draft.pdf", use_container_width=True)
        with e4:
            st.button("🖨️ Print View", use_container_width=True, help="Use browser print (Cmd+P) after clicking")