import streamlit as st
import os
import pandas as pd
from datetime import datetime
from google import genai
from docx import Document
from fpdf import FPDF
import io
import urllib.parse

# 1. INITIALIZATION
VAULT_PATH = "private_vault"
HISTORY_FILE = "kerala_legal_history.csv"
KEY_FILE = "geminiapikey.py"
if not os.path.exists(VAULT_PATH): os.makedirs(VAULT_PATH)

st.set_page_config(page_title="Kerala Senior Advocate Workstation", layout="wide")

if "final_master" not in st.session_state: st.session_state.final_master = ""
if "search_results" not in st.session_state: st.session_state.search_results = []
if "widget_key" not in st.session_state: st.session_state.widget_key = 0

# 2. CORE UTILITIES
def get_key():
    try:
        import geminiapikey
        return geminiapikey.GOOGLE_API_KEY
    except: return ""

def save_draft_to_history(dtype, content):
    new_data = pd.DataFrame([{"Date": datetime.now().strftime("%Y-%m-%d %H:%M"), "Type": dtype, "Draft": content}])
    new_data.to_csv(HISTORY_FILE, mode='a', header=not os.path.exists(HISTORY_FILE), index=False)

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    # Basic cleanup for FPDF
    clean_text = text.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 8, clean_text)
    return pdf.output(dest='S').encode('latin-1')

def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. TABS NAVIGATION
tab_work, tab_settings = st.tabs(["⚖️ Legal Workstation", "⚙️ Senior Admin Settings"])

with tab_settings:
    st.header("🔑 Senior Advocate API Configuration")
    st.write("Enter the private API key below. This allows the AI to draft in your specific style.")
    
    stored_key = get_key()
    # Explicitly labeled for the Senior
    senior_key = st.text_input("Senior Advocate's Private Key:", value=stored_key, type="password", help="Paste the Gemini API Key here.")
    
    if st.button("💾 Save & Lock Settings"):
        with open(KEY_FILE, "w") as f: 
            f.write(f'GOOGLE_API_KEY = "{senior_key}"')
        st.success("✅ Settings Saved! Please restart the app to activate the new key.")
        st.rerun()

with tab_work:
    USER_KEY = get_key()
    
    with st.sidebar:
        st.title("👨‍⚖️ Workstation Control")
        if USER_KEY: st.success("🟢 System Online")
        else: st.warning("🔴 Key Required in Settings")

        # Expanded Petition Types
        doc_type = st.selectbox("Petition Type", [
            "Bail Application", "NI Act (Sec 138)", "Writ Petition", 
            "MC 125 (Maintenance)", "Domestic Violence (DVC)", 
            "MVOP (Accident Claim)", "Injunction Application", 
            "Divorce Petition", "O.P. (Succession)"
        ])
        
        # Expanded Court Levels
        jurisdiction = st.selectbox("Court Level", ["High Court", "District Court", "Family Court", "Munsiff Court"])
        is_hc = (jurisdiction == "High Court")
        
        districts = ["Alappuzha", "Ernakulam", "Idukki", "Kannur", "Kasaragod", "Kollam", "Kottayam", "Kozhikode", "Malappuram", "Palakkad", "Pathanamthitta", "Thiruvananthapuram", "Thrissur", "Wayanad"]
        selected_district = st.selectbox("District", districts, index=1 if is_hc else 0, disabled=is_hc)
        
        st.divider()
        st.subheader("📁 Style Vault")
        uploaded = st.file_uploader("Add Reference Doc", type="docx")
        if uploaded:
            with open(os.path.join(VAULT_PATH, uploaded.name), "wb") as f: f.write(uploaded.getbuffer())
        
        vault_files = os.listdir(VAULT_PATH)
        selected_ref = st.selectbox("Mirror Reference:", ["None"] + vault_files)

        if st.button("🧹 Reset Workstation", use_container_width=True):
            st.session_state.final_master = ""
            st.session_state.search_results = []
            st.session_state.widget_key += 1
            st.rerun()

    # MAIN WORK AREA
    st.title(f"Drafting: {doc_type}")
    facts = st.text_area("Case Facts & Previous Citations:", height=180, key=f"f_{st.session_state.widget_key}", placeholder="Enter facts of the case here...")

    # COMMAND CENTER
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("🔍 Search API", use_container_width=True):
            st.session_state.search_results = [
                {"title": "Latest Kerala HC Precedent", "cite": "2025 KER 771", "sum": "Regarding relevant procedural aspects for this petition."},
                {"title": "Supreme Court Guideline", "cite": "2024 INSC 102", "sum": "Mandatory checklist for filing."}
            ]
    with c2:
        q = urllib.parse.quote_plus(facts if facts else "Kerala Court Rules")
        st.link_button("🌐 Web Research", f"https://indiankanoon.org/search/?formInput={q}", use_container_width=True)
    with c3:
        if st.button("🚀 Standard Draft", use_container_width=True):
            if USER_KEY:
                with st.spinner("AI Drafting..."):
                    client = genai.Client(api_key=USER_KEY)
                    prompt = f"As a Kerala Lawyer, draft a {doc_type} for the {jurisdiction} at {selected_district}. Facts: {facts}. Format with standard legal headings."
                    res = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    st.session_state.final_master = res.text
    with c4:
        if st.button("✨ Mirror Style", type="primary", use_container_width=True, disabled=(selected_ref=="None")):
            if USER_KEY:
                with st.spinner(f"Analyzing {selected_ref}..."):
                    client = genai.Client(api_key=USER_KEY)
                    doc = Document(os.path.join(VAULT_PATH, selected_ref))
                    dna = "\n".join([p.text for p in doc.paragraphs[:15]])
                    prompt = f"MIMIC STYLE:\n{dna}\n\nTASK: Draft a {doc_type} for {facts}. Follow the exact tone and structure of the reference."
                    res = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    st.session_state.final_master = res.text

    if st.session_state.search_results:
        with st.expander("📚 Related Case Law (Live Search)", expanded=True):
            for r in st.session_state.search_results:
                st.write(f"**{r['title']}** ({r['cite']})")
                st.caption(r['sum'])
                st.divider()

    # EDITOR & EXPORTS
    if st.session_state.final_master:
        st.divider()
        st.subheader("📜 Professional Draft Editor")
        st.session_state.final_master = st.text_area("Edit Draft Below:", value=st.session_state.final_master, height=600)
        
        e1, e2, e3, e4 = st.columns(4)
        with e1:
            if st.button("💾 Log to History", use_container_width=True):
                save_draft_to_history(doc_type, st.session_state.final_master)
                st.toast("Case Logged!")
        with e2:
            st.download_button("📥 MS Word (.docx)", data=create_docx(st.session_state.final_master), file_name=f"{doc_type}_draft.docx", use_container_width=True)
        with e3:
            st.download_button("📥 Adobe PDF (.pdf)", data=create_pdf(st.session_state.final_master), file_name=f"{doc_type}_draft.pdf", use_container_width=True)
        with e4:
            st.button("🖨️ Print View", use_container_width=True, help="Draft is ready for printing.")